"""
Word文档生成器 - 将提取的帧和分析结果生成Word文档
支持多种样式和格式
"""

import os
import io
import tempfile

from PIL import Image

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.enum.section import WD_ORIENT
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False


# 预设文档样式主题
THEMES = {
    'professional': {
        'name': '专业商务',
        'title_color': (31, 73, 125),
        'heading_color': (0, 84, 166),
        'text_color': (0, 0, 0),
        'accent_color': (31, 73, 125),
        'font_name': '微软雅黑',
    },
    'modern': {
        'name': '现代简约',
        'title_color': (51, 51, 51),
        'heading_color': (0, 112, 192),
        'text_color': (89, 89, 89),
        'accent_color': (0, 112, 192),
        'font_name': '微软雅黑',
    },
    'academic': {
        'name': '学术论文',
        'title_color': (0, 0, 0),
        'heading_color': (0, 0, 128),
        'text_color': (0, 0, 0),
        'accent_color': (0, 0, 128),
        'font_name': '宋体',
    },
    'clean': {
        'name': '清晰简洁',
        'title_color': (0, 0, 0),
        'heading_color': (0, 0, 0),
        'text_color': (64, 64, 64),
        'accent_color': (100, 100, 100),
        'font_name': '微软雅黑',
    },
}


class WordGenerator:
    def __init__(self, theme='modern', title='视频要点提取', subtitle=''):
        if not _HAS_DOCX:
            raise RuntimeError("请先安装 python-docx: pip install python-docx")
        self.doc = Document()
        self.theme_name = theme
        self.theme = THEMES.get(theme, THEMES['modern'])

        # 设置页面为A4横向（更像PPT）
        section = self.doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)

        # 默认字体
        self._set_default_font()

        self.title_text = title
        self.subtitle_text = subtitle

    def _set_default_font(self):
        """设置默认中文字体"""
        style = self.doc.styles['Normal']
        style.font.name = self.theme['font_name']
        style._element.rPr.rFonts.set(qn('w:eastAsia'), self.theme['font_name'])
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor(*self.theme['text_color'])

    def _rgb(self, rgb_tuple):
        return RGBColor(*rgb_tuple)

    def _add_heading(self, text, level=1):
        """添加标题"""
        p = self.doc.add_heading(text, level=level)
        for run in p.runs:
            if level == 1:
                run.font.color.rgb = self._rgb(self.theme['title_color'])
                run.font.size = Pt(24)
            elif level == 2:
                run.font.color.rgb = self._rgb(self.theme['heading_color'])
                run.font.size = Pt(16)
            else:
                run.font.color.rgb = self._rgb(self.theme['heading_color'])
                run.font.size = Pt(14)
            run.font.bold = True
        return p

    def _add_paragraph(self, text='', bold=False, italic=False, size=11, color=None, align='left'):
        """添加段落"""
        p = self.doc.add_paragraph()
        if align == 'center':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == 'right':
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if text:
            run = p.add_run(text)
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.italic = italic
            if color:
                run.font.color.rgb = self._rgb(color)
            else:
                run.font.color.rgb = self._rgb(self.theme['text_color'])
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.space_before = Pt(3)
        return p

    def _add_bullet_point(self, text, level=0):
        """添加项目符号"""
        p = self.doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = self._rgb(self.theme['text_color'])
        p.paragraph_format.space_after = Pt(4)
        return p

    def _add_divider(self):
        """添加分隔线（用空段落模拟）"""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run('─' * 80)
        run.font.size = Pt(8)
        run.font.color.rgb = self._rgb((200, 200, 200))

    def add_title_page(self):
        """添加封面"""
        # 留空
        self.doc.add_paragraph()

        # 主标题
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(self.title_text)
        run.font.size = Pt(36)
        run.font.bold = True
        run.font.color.rgb = self._rgb(self.theme['title_color'])

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(self.subtitle_text)
        run.font.size = Pt(18)
        run.font.color.rgb = self._rgb(self.theme['accent_color'])

        # 底部信息
        for _ in range(6):
            self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('VideoToPPT · 自动生成')
        run.font.size = Pt(12)
        run.font.color.rgb = self._rgb((128, 128, 128))

    def add_page(self, img_rgb, title, bullet_points=None,
                 include_image=True, image_position='right'):
        """
        添加一页内容（类似PPT的一页）

        参数:
            img_rgb: numpy.ndarray (H, W, 3) RGB
            title: 页面标题
            bullet_points: list[str] 要点列表
            include_image: 是否包含截图
            image_position: 'right' | 'left' | 'below'  图片位置
        """
        # 添加分页
        self._add_divider()

        # 标题
        self._add_heading(title, level=2)

        if include_image:
            if image_position == 'below':
                # 文字在上，图片在下
                if bullet_points:
                    for bp in bullet_points:
                        if bp.strip():
                            self._add_bullet_point(bp.strip())
                    self.doc.add_paragraph()

                self._add_image(img_rgb, width=Inches(6))
            else:
                # 图文并排
                # 创建表格实现左右布局
                table = self.doc.add_table(rows=1, cols=2)
                table.autofit = False
                if image_position == 'right':
                    cell_text = table.cell(0, 0)
                    cell_img = table.cell(0, 1)
                else:
                    cell_text = table.cell(0, 1)
                    cell_img = table.cell(0, 0)

                # 文字单元格
                cell_text.width = Inches(5.5)
                p = cell_text.paragraphs[0]
                if bullet_points:
                    for i, bp in enumerate(bullet_points):
                        if bp.strip():
                            run = p.add_run('• ' + bp.strip())
                            run.font.size = Pt(11)
                            run.font.color.rgb = self._rgb(self.theme['text_color'])
                            if i < len(bullet_points) - 1:
                                p.add_run('\n')
                else:
                    run = p.add_run('（无文本内容）')
                    run.font.italic = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = self._rgb((128, 128, 128))

                # 图片单元格
                cell_img.width = Inches(5.0)
                p = cell_img.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._add_image_to_cell(cell_img, img_rgb, width=Inches(4.5))
        else:
            # 仅文字
            if bullet_points:
                for bp in bullet_points:
                    if bp.strip():
                        self._add_bullet_point(bp.strip())
            else:
                self._add_paragraph('（无文本内容）', italic=True)

        self.doc.add_paragraph()

    def _add_image(self, img_rgb, width=None):
        """添加图片"""
        img_path = self._temp_image_path(img_rgb)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        if width:
            run.add_picture(img_path, width=width)
        else:
            run.add_picture(img_path)
        self._cleanup_temp(img_path)

    def _add_image_to_cell(self, cell, img_rgb, width=None):
        """向Word表格单元格中添加图片"""
        img_path = self._temp_image_path(img_rgb)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        if width:
            run.add_picture(img_path, width=width)
        else:
            run.add_picture(img_path)
        self._cleanup_temp(img_path)

    def _temp_image_path(self, img_rgb):
        """将numpy图像保存为临时文件"""
        pil = Image.fromarray(img_rgb)
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
        tmp.close()
        pil.save(tmp.name, 'PNG')
        return tmp.name

    def _cleanup_temp(self, path):
        try:
            if os.path.exists(path):
                os.unlink(path)
        except Exception:
            pass

    def save(self, output_path):
        """保存Word文档"""
        self.doc.save(output_path)
        return output_path


def build_document(frames_with_analysis, output_path,
                  theme='modern',
                  title='视频要点',
                  subtitle='由 VideoToPPT 自动生成',
                  image_position='right',
                  include_ocr=True,
                  progress_callback=None):
    """
    从分析结果构建完整Word文档

    参数:
        frames_with_analysis: list[dict]
        output_path: 输出文件路径
        theme: 主题名
        image_position: 'right' | 'left' | 'below'
        include_ocr: 是否包含OCR文本
        progress_callback: 进度回调
    """
    gen = WordGenerator(theme=theme, title=title, subtitle=subtitle)
    gen.add_title_page()

    total = len(frames_with_analysis)
    for i, item in enumerate(frames_with_analysis):
        img = item.get('image')
        analysis = item.get('analysis', {})
        title_txt = analysis.get('title') or analysis.get('auto_title') or f"幻灯片 {i + 1}"

        bullets = []
        if include_ocr and analysis.get('bullet_points'):
            bullets = analysis['bullet_points']
        elif include_ocr and analysis.get('content'):
            bullets = [b for b in analysis['content'].split('\n') if b.strip()]

        gen.add_page(
            img_rgb=img,
            title=title_txt,
            bullet_points=bullets,
            include_image=True,
            image_position=image_position,
        )

        if progress_callback:
            percent = int(50 + 50.0 * (i + 1) / max(total, 1))
            progress_callback(percent, f"生成第 {i + 1}/{total} 页...")

    gen.save(output_path)
    return output_path
